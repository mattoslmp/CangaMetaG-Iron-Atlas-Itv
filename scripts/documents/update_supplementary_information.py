#!/usr/bin/env python3
"""Update final Supplementary Information with grouped-only S40, both S67 layouts and Table 16."""
from __future__ import annotations
import argparse, copy, shutil
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt

S40G='SupplementaryFigure40_ST8_external_iron_rich_module_completeness_by_environmental_group'
S67O='SupplementaryFigure67_lagoon_plus_external_iron_rich_module_completeness_KEMET_style_3state_heatmap'
S67G='SupplementaryFigure67_lagoon_plus_external_iron_rich_module_completeness_by_environmental_group'

def clear_paragraph(paragraph):
  for child in list(paragraph._p): paragraph._p.remove(child)

def set_image(paragraph, path: Path):
  clear_paragraph(paragraph)
  paragraph.alignment=1
  paragraph.add_run().add_picture(str(path),width=Inches(15.0))

def insert_after(anchor, elements):
  current=anchor._p
  for element in elements:
    current.addnext(element)
    current=element

def make_image_para(doc: Document, path: Path):
  p=doc.add_paragraph(); p.alignment=1; p.add_run().add_picture(str(path),width=Inches(15.0)); return p

def make_caption_para(doc: Document, text: str):
  p=doc.add_paragraph(); r=p.add_run(text); r.font.size=Pt(9); return p

def remove_unused_image_parts(docx_path: Path) -> None:
  """Remove image relationships/media no longer referenced after targeted deletion."""
  import os, zipfile
  from xml.etree import ElementTree as ET
  ns_r='http://schemas.openxmlformats.org/officeDocument/2006/relationships'
  with zipfile.ZipFile(docx_path,'r') as zin:
    members={name:zin.read(name) for name in zin.namelist()}
  used_by_part={}
  for name,data in members.items():
    if not name.startswith('word/') or not name.endswith('.xml'): continue
    try: root=ET.fromstring(data)
    except ET.ParseError: continue
    used=set()
    for elem in root.iter():
      for attr,val in elem.attrib.items():
        if attr.startswith('{'+ns_r+'}'): used.add(val)
    used_by_part[name]=used
  remove_names=set()
  for rel_name,data in list(members.items()):
    if not rel_name.startswith('word/_rels/') or not rel_name.endswith('.rels'): continue
    source='word/'+rel_name.split('/')[-1][:-5]
    used=used_by_part.get(source,set())
    try: root=ET.fromstring(data)
    except ET.ParseError: continue
    changed=False
    for rel in list(root):
      rid=rel.attrib.get('Id',''); target=rel.attrib.get('Target',''); rtype=rel.attrib.get('Type','')
      if rtype.endswith('/image') and rid not in used:
        root.remove(rel); changed=True
        target_path=os.path.normpath(os.path.join(os.path.dirname(source),target)).replace('\\','/')
        if target_path.startswith('word/media/'): remove_names.add(target_path)
    if changed: members[rel_name]=ET.tostring(root,encoding='utf-8',xml_declaration=True)
  for name in remove_names:
    target_file=os.path.basename(name)
    if not any(rel.endswith('.rels') and target_file.encode() in data for rel,data in members.items()): members.pop(name,None)
  tmp=docx_path.with_suffix('.cleaning.tmp.docx')
  with zipfile.ZipFile(tmp,'w',zipfile.ZIP_DEFLATED) as zout:
    for name,data in members.items(): zout.writestr(name,data)
  tmp.replace(docx_path)

def main()->int:
  ap=argparse.ArgumentParser(description=__doc__)
  ap.add_argument('--app-root',type=Path,required=True)
  ap.add_argument('--article-root',type=Path,required=True)
  a=ap.parse_args(); app=a.app_root.resolve(); art=a.article_root.resolve()
  final_dir=art/'01_Manuscript'/'FINAL_SUBMISSION_FILES'
  source=art/'01_Manuscript'/'SOURCE_REVISION4'/'Supplementary_Information_ISME_REVISION4_ORDINATION_CORRECTED.docx'
  table_doc=Document(str(app/'tables'/'Supplementary_Table_16_final_scripts.docx'))
  doc=Document(str(source))
  paragraphs=doc.paragraphs
  # Replace canonical S40 slots with the environmental-group version only.
  replacements={108:(S40G,1),110:(S40G,2),171:(S67O,1),173:(S67O,2)}
  for idx,(stem,panel) in replacements.items():
    set_image(paragraphs[idx],app/'outputs'/'final_publication_figures'/f'{stem}_P{panel:03d}.png')
  paragraphs[109].text='Supplementary Figure 40. Final environmental-group layout of the thematic three-state KEGG/KEMET module-completeness heatmap for curated external iron-rich records. All records are retained and placed side by side by environmental group; module rows, cell values, completeness states, colours, axis orientation and record identities are unchanged relative to the immutable source matrix. Short record labels correspond to Supplementary Table 15. Panel P001 of P002.'
  paragraphs[111].text='Supplementary Figure 40. Final environmental-group layout of the thematic three-state KEGG/KEMET module-completeness heatmap for curated external iron-rich records. All records are retained and placed side by side by environmental group; module rows, cell values, completeness states, colours, axis orientation and record identities are unchanged relative to the immutable source matrix. Short record labels correspond to Supplementary Table 15. Panel P002 of P002.'
  # Delete only the former duplicate S40 alternative image/caption pairs.
  for idx in (116,115,114,113):
    element=paragraphs[idx]._p; element.getparent().remove(element)
  paragraphs=doc.paragraphs
  # Insert grouped S67 immediately after the original S67 panel 2 caption.
  anchor=paragraphs[170]
  # Do not add explicit page-break-only paragraphs: the full-page images flow to
  # the next page naturally. Explicit breaks before these tall panels created blank
  # pages in LibreOffice/Word pagination.
  img1=make_image_para(doc,app/'outputs'/'final_publication_figures'/f'{S67G}_P001.png')
  cap1=make_caption_para(doc,'Alternative environmental-group layout — Supplementary Figure 67, Panel P001 of P002. The identical combined Amazonian-plus-external S67 matrix is shown with only the column order changed so that samples and external records from the same environmental group are adjacent. All modules, values, states, colors and axis orientation are unchanged.')
  img2=make_image_para(doc,app/'outputs'/'final_publication_figures'/f'{S67G}_P002.png')
  cap2=make_caption_para(doc,'Alternative environmental-group layout — Supplementary Figure 67, Panel P002 of P002. The identical combined Amazonian-plus-external S67 matrix is shown with only the column order changed so that samples and external records from the same environmental group are adjacent. All modules, values, states, colors and axis orientation are unchanged.')
  # Newly appended temporary elements must be detached before relocation.
  elements=[img1._p,cap1._p,img2._p,cap2._p]
  for e in elements:
    parent=e.getparent()
    if parent is not None: parent.remove(e)
  insert_after(anchor,elements)
  # Replace the old abbreviated Table 16 with the full editable table.
  if len(doc.tables)<2: raise RuntimeError('Expected existing Supplementary Table 16')
  old=doc.tables[1]._tbl; old.getparent().remove(old)
  caption=next(p for p in doc.paragraphs if p.text.strip().startswith('Supplementary Table 16.'))
  new_tbl=copy.deepcopy(table_doc.tables[0]._tbl)
  caption._p.addnext(new_tbl)
  output=final_dir/'Supplementary_Information_ISME_REPRODUCIBILITY_FINAL.docx'
  doc.save(str(output))
  # Re-open and save through python-docx to normalize the OPC package. This keeps
  # the document editable and avoids unsafe low-level deletion of media relations.
  normalized=output.with_suffix('.normalized.tmp.docx')
  Document(str(output)).save(str(normalized))
  normalized.replace(output)
  shutil.copy2(output, final_dir/'Supplementary_Information_ISME_REVISION4_ORDINATION_CORRECTED.docx')
  print(f'PASS: {output}')
  return 0
if __name__=='__main__': raise SystemExit(main())
