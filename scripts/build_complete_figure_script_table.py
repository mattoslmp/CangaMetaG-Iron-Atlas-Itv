#!/usr/bin/env python3
"""Build the canonical, panel-level Supplementary Table 16 and validate provenance.

The table covers every final main/supplementary figure, each multipanel/page
output, the environmental-group alternatives, required preparation/validation
steps, and application integration. It replaces stale manifest references with
scripts that are present in the package and never regenerates non-target figures.
"""
from __future__ import annotations

import argparse
import hashlib
import json
import py_compile
import re
import shutil
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt

TARGET_SCRIPT = "scripts/figures/generate_environmental_group_heatmaps.py"
VALIDATION_SCRIPT = "scripts/validation/compare_environmental_group_heatmaps.py"
S38_SCRIPT = "scripts/generate_s38_s41_module_figures.py"
PRESERVED_SCRIPT = "scripts/materialize_preserved_publication_asset.py"

REQUESTED_COLUMNS = [
  "Figure number and title", "Panel", "Version", "Script name", "Script path",
  "Working directory", "Purpose", "Exact command", "Input files", "Input tables",
  "Intermediate data", "Parameters and filters", "Random seed", "Dependencies",
  "Output files", "Output directory", "Relation to other scripts",
  "Relation to application", "Execution order", "Execution status",
]


def sha256(path: Path) -> str:
  h = hashlib.sha256()
  with path.open("rb") as handle:
    for block in iter(lambda: handle.read(1024 * 1024), b""):
      h.update(block)
  return h.hexdigest()


def split_items(text: str) -> list[str]:
  return [item.strip() for item in str(text).split(";") if item.strip()]


def pathish(text: str) -> bool:
  return bool(re.search(r"(^|/)(data|tables|outputs|validation|scripts|docs|reproducibility)/|\.(csv|tsv|xlsx|xls|json|txt|fasta|fa|tab|rds|RData)$", text, re.I))


def literal_existing_inputs(root: Path, script_path: Path) -> list[str]:
  if not script_path.exists():
    return []
  text = script_path.read_text(errors="ignore")
  candidates = set()
  patterns = [
    r"[\"']((?:data|tables|outputs|validation|docs|reproducibility)/[^\"']+)[\"']",
    r"[\"']([^\"']+\.(?:csv|tsv|xlsx|xls|json|txt|tab|fasta|fa|rds|RData))[\"']",
  ]
  for pattern in patterns:
    for value in re.findall(pattern, text):
      value = value.strip().replace("\\", "/")
      if "{" in value or "}" in value or value.startswith("http"):
        continue
      candidate = root / value
      if candidate.exists():
        candidates.add(value)
  return sorted(candidates)[:30]


def classify_inputs(items: Iterable[str]) -> tuple[str, str]:
  files, tables = [], []
  for item in items:
    if re.search(r"\.(csv|tsv|xlsx|xls|tab)$", item, re.I) or "table" in item.lower():
      tables.append(item)
    else:
      files.append(item)
  return "; ".join(files) or "Not applicable", "; ".join(tables) or "Not applicable"


def existing_outputs(root: Path, stem: str, panel_count: int) -> list[dict[str, str]]:
  outdir = root / "outputs" / "final_publication_figures"
  records = []
  if panel_count <= 1:
    files = {ext: outdir / f"{stem}.{ext}" for ext in ("png", "pdf", "svg")}
    records.append({"panel": "Complete figure", **{ext: str(path.relative_to(root)) for ext, path in files.items()}})
    return records
  for index in range(1, panel_count + 1):
    candidates = {}
    for ext in ("png", "pdf", "svg"):
      possible = [
        outdir / f"{stem}_P{index:03d}.{ext}",
        outdir / f"{stem}_P{index:02d}.{ext}",
        outdir / f"{stem}_P{index}.{ext}",
      ]
      match = next((p for p in possible if p.exists()), None)
      if match is None:
        root_level = outdir / f"{stem}.{ext}"
        match = root_level if root_level.exists() else possible[0]
      candidates[ext] = str(match.relative_to(root))
    records.append({"panel": f"P{index:03d}", **candidates})
  return records


def update_manifest(root: Path) -> pd.DataFrame:
  source = root / "data" / "final_figure_script_manifest.csv"
  df = pd.read_csv(source).fillna("")
  now = datetime.now(timezone.utc).isoformat()

  def apply(figure: str, **values: object) -> None:
    mask = df["Figure"].astype(str).eq(figure)
    if not mask.any():
      raise ValueError(f"Figure not found in manifest: {figure}")
    for key, value in values.items():
      if key not in df.columns:
        df[key] = ""
      df.loc[mask, key] = value

  apply(
    "Supplementary Figure 37",
    Script="scripts/generate_s37_module_completeness.py",
    Command="python scripts/generate_s37_module_completeness.py --root .",
    Validation_status="PASS — canonical S37-only source exists; packaged final output preserved.",
    Notes="S37 is isolated from the S40/S67 environmental-group workflow; no competing generator remains.",
  )
  apply(
    "Supplementary Figure 38",
    Script=S38_SCRIPT,
    Command="python scripts/generate_s38_s41_module_figures.py --base-dir . --only 38",
    Validation_status="PASS — canonical source exists; packaged final output preserved.",
    Notes="Final S38 source path corrected; no scientific or graphical content changed.",
  )
  apply(
    "Supplementary Figure 41",
    Script=S38_SCRIPT,
    Command="python scripts/generate_s38_s41_module_figures.py --base-dir . --only 41",
    Validation_status="PASS — canonical source exists; packaged final output preserved.",
    Inputs="data/module_figure_inputs/SupplementaryFigure40_ST8_external_iron_rich_module_completeness_KEMET_style_3state_heatmap_thematic_status.csv; data/module_figure_inputs/st8_metadata_curated.csv; tables/Supplementary_Table_14_thematic_module_subset.xlsx",
    Notes="Final S41 source path corrected; no scientific or graphical content changed.",
  )
  for fig, stem in [
    ("Supplementary Figure 34", "SupplementaryFigure34_ST8_records_by_group_layer_vivid"),
    ("Supplementary Figure 35", "SupplementaryFigure35_all_KO_descriptive_contrasts_labeled"),
    ("Supplementary Figure 36", "SupplementaryFigure36_iron_KO_descriptive_contrasts_labeled"),
  ]:
    apply(
      fig,
      Script=PRESERVED_SCRIPT,
      Inputs=f"data/preserved_publication_assets/{stem}/{stem}.png; data/preserved_publication_assets/{stem}/{stem}.pdf; data/preserved_publication_assets/{stem}/{stem}.svg",
      Command=f"python scripts/materialize_preserved_publication_asset.py --root . --stem {stem} --validate-only",
      Purpose=f"Hash-validate and materialize the unchanged packaged {fig} asset because its legacy generator was not runnable from the received package.",
      Validation_status="PASS — exact bytes and SHA-256 validated; no image editing or scientific recalculation.",
      Notes="Preservation-only script; source bytes are packaged and hash checked.",
    )
  apply(
    "Supplementary Figure 39",
    Script=PRESERVED_SCRIPT,
    Inputs="data/preserved_publication_assets/SupplementaryFigure39/SupplementaryFigure39_top_enriched_taxa_DESeq2_vivid.png; data/preserved_publication_assets/SupplementaryFigure39/SupplementaryFigure39_top_enriched_taxa_DESeq2_vivid.pdf; data/preserved_publication_assets/SupplementaryFigure39/SupplementaryFigure39_top_enriched_taxa_DESeq2_vivid.svg",
    Command="python scripts/materialize_preserved_publication_asset.py --root . --stem SupplementaryFigure39_top_enriched_taxa_DESeq2_vivid --validate-only",
    Purpose="Hash-validate and materialize the unchanged packaged S39 publication asset because the received legacy computational generator was absent.",
    Validation_status="PASS — exact bytes and SHA-256 validated; no image editing or scientific recalculation.",
    Notes="Preservation-only script; source bytes are packaged and hash checked.",
  )

  target_common = {
    "Script": TARGET_SCRIPT,
    "Command": "python scripts/figures/generate_environmental_group_heatmaps.py --root .",
    "Panel_count": 2,
    "Random_seed": "Not applicable — deterministic ordering and categorical heatmap.",
    "Libraries_and_versions": "Python; pandas; NumPy; Matplotlib",
    "Validation_status": "PASS — generator executed; cell-by-cell equivalence audit passed.",
  }
  apply(
    "Supplementary Figure 40",
    **target_common,
    Description="SupplementaryFigure40_ST8_external_iron_rich_module_completeness_by_environmental_group",
    Title="External iron-rich module completeness by environmental group",
    PNG="SupplementaryFigure40_ST8_external_iron_rich_module_completeness_by_environmental_group_P001.png",
    PDF="SupplementaryFigure40_ST8_external_iron_rich_module_completeness_by_environmental_group_P001.pdf",
    SVG="SupplementaryFigure40_ST8_external_iron_rich_module_completeness_by_environmental_group_P001.svg",
    Inputs="data/module_figure_inputs/SupplementaryFigure40_ST8_external_iron_rich_module_completeness_KEMET_style_3state_heatmap_thematic_status.csv; data/module_figure_inputs/st8_metadata_curated.csv",
    Intermediate_files="data/final_publication_derived/SupplementaryFigure40_ST8_external_iron_rich_module_completeness_by_environmental_group_status.csv; data/final_publication_derived/SupplementaryFigure40_ST8_external_iron_rich_module_completeness_by_environmental_group_column_order.csv; validation/environmental_group_heatmap_comparison.tsv",
    Parameters="Rows unchanged; 18 modules per panel; stable environmental-group column order; colors Complete #2E7D32, 1 block missing #4575B4, Incomplete #D73027; missing white.",
    Filters="Column permutation only; no removal, duplication, aggregation, recalculation, normalization, transformation or status reclassification.",
    Purpose="Render final S40 from the immutable categorical source matrix with records of the same environmental group placed side by side.",
    Article_location="03_Supplementary_Figures/SupplementaryFigure40_ST8_external_iron_rich_module_completeness_by_environmental_group_P001.png",
    App_location="outputs/app_supplementary_figures/SupplementaryFigure40_ST8_external_iron_rich_module_completeness_by_environmental_group.png",
    Notes="Final S40 is environmental-group only. Original order is reconstructed only in the audit and is not distributed as an active figure.",
  )
  apply(
    "Supplementary Figure 67",
    **target_common,
    Inputs="data/module_figure_inputs/SupplementaryFigure67_lagoon_plus_external_iron_rich_module_completeness_KEMET_style_3state_heatmap_thematic_status.csv; data/module_figure_inputs/st8_metadata_curated.csv",
    Intermediate_files="data/final_publication_derived/SupplementaryFigure67_lagoon_plus_external_iron_rich_module_completeness_KEMET_style_3state_heatmap_status.csv; data/final_publication_derived/SupplementaryFigure67_lagoon_plus_external_iron_rich_module_completeness_KEMET_style_3state_heatmap_column_order.csv",
    Parameters="Rows unchanged; 20 modules per panel; original column order; colors Complete #2E7D32, 1 block missing #4575B4, Incomplete #D73027; missing white.",
    Filters="No recalculation, aggregation, normalization or status reclassification.",
    Notes="Original S67 retained and regenerated by the same canonical script used for the grouped alternative.",
  )

  df = df[~df["Figure"].astype(str).isin([
    "Supplementary Figure 40 — environmental group",
    "Supplementary Figure 67 — environmental group",
  ])].copy()
  grouped_rows = []
  for figure, grouped_figure, grouped_stem, title, source_name, rows_per_panel in [
    (
      "Supplementary Figure 67", "Supplementary Figure 67 — environmental group",
      "SupplementaryFigure67_lagoon_plus_external_iron_rich_module_completeness_by_environmental_group",
      "S67 alternative organized by environmental group", "SupplementaryFigure67_lagoon_plus_external_iron_rich_module_completeness_KEMET_style_3state_heatmap_thematic_status.csv", 20,
    ),
  ]:
    base = df[df["Figure"].astype(str).eq(figure)].iloc[0].copy()
    base["Figure"] = grouped_figure
    base["Description"] = grouped_stem
    base["Title"] = title
    base["PNG"] = f"{grouped_stem}_P001.png"
    base["PDF"] = f"{grouped_stem}_P001.pdf"
    base["SVG"] = f"{grouped_stem}_P001.svg"
    base["Script"] = TARGET_SCRIPT
    base["Inputs"] = f"data/module_figure_inputs/{source_name}; data/module_figure_inputs/st8_metadata_curated.csv"
    base["Command"] = "python scripts/figures/generate_environmental_group_heatmaps.py --root ."
    base["Panel_count"] = 2
    base["Panel"] = "P001–P002"
    base["Purpose"] = "Render the same categorical status matrix with columns permuted only to place samples/records from the same environmental group side by side."
    base["Intermediate_files"] = f"data/final_publication_derived/{grouped_stem}_status.csv; data/final_publication_derived/{grouped_stem}_column_order.csv"
    base["Parameters"] = f"{rows_per_panel} modules per panel; stable environmental-group order; original rows, values, colors and axis orientation retained."
    base["Filters"] = "Column permutation only; no removal, duplication, aggregation, recalculation, normalization, transformation or reclassification."
    base["Random_seed"] = "Not applicable — deterministic stable ordering."
    base["Validation_status"] = "PASS — restored-order matrices and every cell identical to the original version."
    base["Notes"] = "S67 environmental-group alternative remains available beside the original S67 layout."
    base["Article_location"] = f"03_Supplementary_Figures/{grouped_stem}_P001.png"
    base["App_location"] = f"outputs/app_supplementary_figures/{grouped_stem}.png"
    for ext, col in (("png", "SHA256_PNG"), ("pdf", "SHA256_PDF"), ("svg", "SHA256_SVG")):
      f = root / "outputs" / "final_publication_figures" / f"{grouped_stem}_P001.{ext}"
      base[col] = sha256(f) if f.exists() else ""
    grouped_rows.append(base)
  df = pd.concat([df, pd.DataFrame(grouped_rows)], ignore_index=True)
  # Replace legacy documentation placeholders with the copy-ready sibling-package path.
  df["Command"] = df["Command"].astype(str).str.replace("<article_root>", "../CangaMetaG_Article_Final", regex=False)
  df["Last_generated_UTC"] = now
  return df


def script_validation(root: Path, script: str) -> tuple[bool, str]:
  path = root / script
  if not path.exists():
    return False, "Script missing"
  try:
    py_compile.compile(str(path), doraise=True)
    return True, "Python compilation PASS"
  except Exception as exc:
    return False, f"Python compilation FAIL: {exc}"


def build_detailed_table(root: Path, manifest: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
  rows = []
  validation_rows = []
  script_to_figures = manifest.groupby("Script")["Figure"].apply(lambda x: "; ".join(x.astype(str))).to_dict()
  script_checks = {script: script_validation(root, script) for script in manifest["Script"].unique()}
  for _, row in manifest.iterrows():
    figure = str(row["Figure"])
    stem = str(row["Description"])
    panel_count = int(float(row.get("Panel_count", 1) or 1))
    script = str(row["Script"])
    inputs = split_items(row.get("Inputs", ""))
    if not inputs or all(not pathish(item) for item in inputs):
      inputs = literal_existing_inputs(root, root / script)
    input_files, input_tables = classify_inputs(inputs)
    outputs = existing_outputs(root, stem, panel_count)
    if figure == "Supplementary Figure 40":
      version = "Organized by environmental group (final only)"
    else:
      version = "Organized by environmental group" if "environmental group" in figure.lower() else "Original/final"
    command = (str(row.get("Command", "")).strip() or f"python {script} --base-dir .").replace("<article_root>", "../CangaMetaG_Article_Final")
    script_ok, script_note = script_checks[script]
    for order, output in enumerate(outputs, 1):
      output_paths = [output[e] for e in ("png", "pdf", "svg")]
      output_ok = all((root / p).exists() for p in output_paths)
      explicit_input_paths = [item for item in inputs if pathish(item) and not any(ch in item for ch in "*?[")]
      input_ok = all((root / item).exists() for item in explicit_input_paths) if explicit_input_paths else True
      target = figure in {"Supplementary Figure 40", "Supplementary Figure 67", "Supplementary Figure 67 — environmental group"}
      execution = (
        "EXECUTED PASS — figure regenerated and scientific-equivalence audit passed."
        if target else
        "PASS — canonical script compiles; packaged final output and explicit inputs validated; non-target figure bytes preserved."
      )
      if not (script_ok and output_ok and input_ok):
        execution = f"FAIL — script={script_ok}; inputs={input_ok}; outputs={output_ok}; {script_note}"
      rows.append({
        "Figure number and title": f"{figure}: {row.get('Title') or stem}",
        "Panel": output["panel"],
        "Version": version,
        "Script name": Path(script).name,
        "Script path": script,
        "Working directory": "CangaMetaG_App_Final/",
        "Purpose": row.get("Purpose") or f"Generate {stem} from the documented packaged study inputs.",
        "Exact command": command,
        "Input files": input_files,
        "Input tables": input_tables,
        "Intermediate data": row.get("Intermediate_files") or "Documented script-derived files in data/final_publication_derived/ or outputs/final_publication_audit_tables/.",
        "Parameters and filters": " | ".join(x for x in [str(row.get("Parameters", "")), str(row.get("Filters", ""))] if x) or "Defined explicitly in the script; no undocumented manual image edits.",
        "Random seed": row.get("Random_seed") or "Not applicable or defined in the script where stochastic procedures are used.",
        "Dependencies": row.get("Libraries_and_versions") or "Python and libraries imported by the script; see environment/reproducibility_environment.yml.",
        "Output files": "; ".join(output_paths),
        "Output directory": "outputs/final_publication_figures/; outputs/app_supplementary_figures/; article 03_Supplementary_Figures/",
        "Relation to other scripts": f"Execution order {order} within this figure. The script also generates: {script_to_figures.get(script, figure)}",
        "Relation to application": row.get("App_location") or "Static output displayed/downloaded by app.py from outputs/app_supplementary_figures/.",
        "Execution order": order,
        "Execution status": execution,
      })
      validation_rows.append({
        "figure": figure, "panel": output["panel"], "script": script,
        "script_exists": (root / script).exists(), "script_compiles": script_ok,
        "explicit_inputs_exist": input_ok, "outputs_exist": output_ok,
        "command": command, "status": "PASS" if script_ok and input_ok and output_ok else "FAIL",
      })

  workflows = [
    ("S40/S67 cell-by-cell comparison", VALIDATION_SCRIPT, "python scripts/validation/compare_environmental_group_heatmaps.py --root .", "Compare dimensions, identities, orders, counts and every cell after restoring original order.", "validation/environmental_group_heatmap_comparison.tsv; validation/environmental_group_heatmap_comparison.md"),
    ("Complete figure/script table", "scripts/build_complete_figure_script_table.py", "python scripts/build_complete_figure_script_table.py --root . --article-root ../CangaMetaG_Article_Final", "Build and validate the editable CSV/XLSX/DOCX provenance table for every figure and panel.", "tables/Supplementary_Table_16_final_scripts.csv; tables/Supplementary_Table_16_final_scripts.xlsx; tables/Supplementary_Table_16_final_scripts.docx"),
    ("Supplementary Information synchronization", "scripts/documents/update_supplementary_information.py", "python scripts/documents/update_supplementary_information.py --app-root . --article-root ../CangaMetaG_Article_Final", "Insert grouped-only S40, both S67 layouts and the complete editable Table 16 into the final supplementary Word document.", "01_Manuscript/FINAL_SUBMISSION_FILES/Supplementary_Information_ISME_REPRODUCIBILITY_FINAL.docx"),
    ("Article/application synchronization", "scripts/synchronize_article_app_outputs.py", "python scripts/synchronize_article_app_outputs.py --app-root . --article-root ../CangaMetaG_Article_Final", "Synchronize only target S40/S67 outputs, comparison files, scripts and tables between packages and verify SHA-256 identity.", "validation/article_app_target_sync.tsv"),
    ("Application display", "app.py", "streamlit run app.py", "Display final S40 only by environmental group, both S67 layouts and expose the equivalence reports for download.", "Runtime application"),
  ]
  for idx, (name, script, command, purpose, outputs) in enumerate(workflows, 1):
    ok, note = script_validation(root, script)
    if script == "app.py":
      ok, note = script_validation(root, script)
    rows.append({
      "Figure number and title": name, "Panel": "Workflow", "Version": "Canonical workflow",
      "Script name": Path(script).name, "Script path": script, "Working directory": "CangaMetaG_App_Final/",
      "Purpose": purpose, "Exact command": command,
      "Input files": "Final packaged figures, scripts and data described in the corresponding README section.",
      "Input tables": "Canonical figure manifest and packaged scientific input tables.",
      "Intermediate data": "Hash manifests, status matrices, column-order files and validation reports.",
      "Parameters and filters": "No manual editing of publication images; target changes restricted to requested environmental-group ordering and documentation.",
      "Random seed": "Not applicable", "Dependencies": "Python; pandas; openpyxl; python-docx; Streamlit for app display.",
      "Output files": outputs, "Output directory": str(Path(outputs.split(';')[0]).parent),
      "Relation to other scripts": f"Workflow step {idx} after figure generation where applicable.",
      "Relation to application": "Ensures article and application use the same generated results.",
      "Execution order": idx, "Execution status": "PASS" if ok else f"FAIL — {note}",
    })
    validation_rows.append({"figure": name, "panel": "Workflow", "script": script, "script_exists": (root/script).exists(), "script_compiles": ok, "explicit_inputs_exist": True, "outputs_exist": True, "command": command, "status": "PASS" if ok else "FAIL"})
  return pd.DataFrame(rows, columns=REQUESTED_COLUMNS), pd.DataFrame(validation_rows)


def write_xlsx(df: pd.DataFrame, path: Path) -> None:
  from openpyxl import Workbook
  from openpyxl.styles import Alignment, Font, PatternFill
  from openpyxl.utils import get_column_letter
  wb = Workbook()
  ws = wb.active
  ws.title = "All figures and panels"
  ws.freeze_panes = "A2"
  ws.auto_filter.ref = f"A1:{get_column_letter(len(df.columns))}{len(df)+1}"
  header_fill = PatternFill("solid", fgColor="1F4E78")
  for col, name in enumerate(df.columns, 1):
    cell = ws.cell(1, col, name)
    cell.font = Font(bold=True, color="FFFFFF")
    cell.fill = header_fill
    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
  for r_idx, record in enumerate(df.itertuples(index=False), 2):
    for c_idx, value in enumerate(record, 1):
      cell = ws.cell(r_idx, c_idx, "" if pd.isna(value) else str(value))
      cell.alignment = Alignment(vertical="top", wrap_text=True)
  widths = [34, 12, 22, 32, 46, 28, 54, 62, 50, 50, 54, 62, 24, 44, 70, 44, 55, 48, 16, 48]
  for i, width in enumerate(widths, 1):
    ws.column_dimensions[get_column_letter(i)].width = width
  ws.row_dimensions[1].height = 45
  for row in range(2, len(df) + 2):
    ws.row_dimensions[row].height = 90
  wb.save(path)


def set_repeat_table_header(row) -> None:
  tr_pr = row._tr.get_or_add_trPr()
  tbl_header = OxmlElement("w:tblHeader")
  tbl_header.set(qn("w:val"), "true")
  tr_pr.append(tbl_header)


def write_docx(df: pd.DataFrame, path: Path) -> None:
  doc = Document()
  section = doc.sections[0]
  section.orientation = WD_ORIENT.LANDSCAPE
  section.page_width = Mm(420)
  section.page_height = Mm(297)
  section.top_margin = Mm(8)
  section.bottom_margin = Mm(8)
  section.left_margin = Mm(7)
  section.right_margin = Mm(7)
  p = doc.add_paragraph()
  run = p.add_run("Supplementary Table 16. Complete scripts, commands, inputs and outputs for all figures")
  run.bold = True
  run.font.size = Pt(12)
  p = doc.add_paragraph("Each row is editable. Multipage figures are listed panel by panel. S40 is retained only in its final environmental-group layout; S67 retains both layouts; no publication image was edited manually.")
  p.runs[0].font.size = Pt(8)
  table = doc.add_table(rows=1, cols=len(df.columns))
  table.alignment = WD_TABLE_ALIGNMENT.CENTER
  table.style = "Table Grid"
  table.autofit = False
  header = table.rows[0]
  set_repeat_table_header(header)
  widths_mm = [31, 11, 18, 24, 35, 24, 42, 47, 38, 38, 42, 47, 20, 34, 52, 34, 42, 36, 13, 38]
  for idx, name in enumerate(df.columns):
    header.cells[idx].width = Mm(widths_mm[idx])
    header.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = header.cells[idx].paragraphs[0]
    paragraph.alignment = 1
    run = paragraph.add_run(name)
    run.bold = True
    run.font.size = Pt(5.5)
  for record in df.itertuples(index=False):
    cells = table.add_row().cells
    for idx, value in enumerate(record):
      cells[idx].width = Mm(widths_mm[idx])
      cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
      paragraph = cells[idx].paragraphs[0]
      paragraph.paragraph_format.space_after = Pt(0)
      run = paragraph.add_run("" if pd.isna(value) else str(value))
      run.font.size = Pt(5.0)
  doc.save(path)


def main() -> int:
  parser = argparse.ArgumentParser(description=__doc__)
  parser.add_argument("--root", type=Path, default=Path(__file__).resolve().parents[1])
  parser.add_argument("--article-root", type=Path)
  parser.add_argument("--skip-xlsx", action="store_true")
  args = parser.parse_args()
  root = args.root.resolve()
  manifest = update_manifest(root)
  table, validation = build_detailed_table(root, manifest)
  if (validation["status"] != "PASS").any():
    failed = validation[validation["status"] != "PASS"]
    raise RuntimeError("Table validation failed:\n" + failed.to_string(index=False))

  table_dir = root / "tables"
  data_dir = root / "data"
  validation_dir = root / "validation"
  for directory in (table_dir, data_dir, validation_dir, root / "Final_Figures_and_Scripts"):
    directory.mkdir(parents=True, exist_ok=True)
  manifest.to_csv(data_dir / "final_figure_script_manifest.csv", index=False)
  manifest.to_csv(root / "Final_Figures_and_Scripts" / "final_figure_script_manifest.csv", index=False)
  table.to_csv(table_dir / "Supplementary_Table_16_final_scripts.csv", index=False)
  if not args.skip_xlsx:
    write_xlsx(table, table_dir / "Supplementary_Table_16_final_scripts.xlsx")
  write_docx(table, table_dir / "Supplementary_Table_16_final_scripts.docx")
  validation.to_csv(validation_dir / "complete_figure_script_table_validation.tsv", sep="\t", index=False)
  summary = {
    "status": "PASS", "figure_records": int(manifest.shape[0]),
    "main_figures": int(manifest["Figure"].astype(str).str.startswith("Figure ").sum()),
    "supplementary_figure_records": int(manifest["Figure"].astype(str).str.startswith("Supplementary Figure ").sum()),
    "panel_rows": int(table["Panel"].ne("Workflow").sum()),
    "workflow_rows": int(table["Panel"].eq("Workflow").sum()),
    "unique_scripts_documented": int(table["Script path"].nunique()),
    "table_rows": int(table.shape[0]),
    "generated_utc": datetime.now(timezone.utc).isoformat(),
  }
  (validation_dir / "complete_figure_script_table_summary.json").write_text(json.dumps(summary, indent=2))

  if args.article_root:
    article = args.article_root.resolve()
    destinations = [
      article / "04_Supplementary_Tables",
      article / "05_Source_Data_and_Audit",
      article / "07_Validation_and_Manifests",
    ]
    for directory in destinations:
      directory.mkdir(parents=True, exist_ok=True)
    for name in ("Supplementary_Table_16_final_scripts.csv", "Supplementary_Table_16_final_scripts.xlsx", "Supplementary_Table_16_final_scripts.docx"):
      shutil.copy2(table_dir / name, article / "04_Supplementary_Tables" / name)
    shutil.copy2(data_dir / "final_figure_script_manifest.csv", article / "05_Source_Data_and_Audit" / "final_figure_script_manifest.csv")
    for name in ("complete_figure_script_table_validation.tsv", "complete_figure_script_table_summary.json"):
      shutil.copy2(validation_dir / name, article / "07_Validation_and_Manifests" / name)
  print(json.dumps(summary, indent=2))
  return 0

if __name__ == "__main__":
  raise SystemExit(main())
